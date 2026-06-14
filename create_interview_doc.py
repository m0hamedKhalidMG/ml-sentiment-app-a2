from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """Helper to set cell borders (optional styling)."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = docx.oxml.OxmlElement(tag)
                tcPr.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "auto")

def add_heading(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80) if level == 1 else RGBColor(0x00, 0x00, 0x00)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(14)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(12)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(12)
    return p

def main():
    doc = Document()
    # Title
    title = doc.add_heading(level=0)
    run = title.add_run('شرح مشروع ML Sentiment Analysis — End-to-End MLOps Pipeline')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_paragraph(doc, 'ده دليل شامل لمشروع "تحليل المشاعر (Sentiment Analysis)" باستخدام تقنيات MLOps و LLMs، مكتوب بالعربي لمساعدتك في الInterview. هشرح كل جزء في المشروع: إحنا عملنا إيه، ليه عملناه، واستخدمنا إيه، وإزاي بيشتغل.')

    # 1. Overview
    add_heading(doc, '1. نظرة عامة على المشروع (Project Overview)', level=1)
    add_paragraph(doc, 'المشروع عبارة عن نظام كامل لتحليل المشاعر (Sentiment Analysis) باستخدام Large Language Model (LLM) اسمه Qwen2.5-1.5B. الفكرة مش بس إننا نبني model، لكن نمر بكل مراحل الـ MLOps Lifecycle:')
    add_bullet(doc, 'Experimentation: نجرب استراتيجيات Few-Shot Prompting ونسجل النتائج.')
    add_bullet(doc, 'Optimization: نضغط الـ Model باستخدام Quantization (GPTQ 4-bit) عشان يشتغل أسرع وأخف.')
    add_bullet(doc, 'Serving: نوفر API للـ Model باستخدام FastAPI + vLLM، ونعمل A/B Testing بين النسخة الأصلية والمضغوطة.')
    add_bullet(doc, 'Monitoring: نراقب أداء الـ Model في الـ Production ونكتشف لو البيانات اتغيرت (Data Drift).')
    add_bullet(doc, 'Deployment: ننشر بـ Kubernetes Canary Deployment عشان ننزل updates بأمان.')
    add_paragraph(doc, 'ليه عملنا كده؟ عشان في الـ Production مش بس الموديل لازم يكون دقيق، لازم يكون سريع، أخف، قابل للمراقبة، والـ Deployment يكون بأمان بدون ما يوقف الخدمة.')

    # 2. Architecture
    add_heading(doc, '2. الـ Architecture (التصميم العام)', level=1)
    add_paragraph(doc, 'التصميم العام للمشروع بيتكون من المراحل دي بالترتيب:')
    add_numbered(doc, 'Training & Few-Shot Experiments: نجرب 1-shot، 3-shot، و 5-shot prompting، ونستخدم MLflow لـ Tracking.')
    add_numbered(doc, 'MLflow Registry: بنسجل أفضل نسخة (Champion) ونسخة التحدي (Challenger) في الـ Model Registry.')
    add_numbered(doc, 'Quantization (GPTQ 4-bit): نضغط الموديل من ~3GB لـ ~1GB.')
    add_numbered(doc, 'FastAPI Traffic Router: بيوزع الـ Traffic (80% V1 و 20% V2) ويسجل الـ Predictions في JSONL.')
    add_numbered(doc, 'vLLM Servers: V1 (FP16) و V2 (GPTQ) بيستقبلوا الـ Requests ويولدوا الـ Responses.')
    add_numbered(doc, 'Evidently Drift Detection + Prometheus: بيكتشفوا لو في Drift في البيانات ويقدروا يبعتوا Alerts.')
    add_numbered(doc, 'Argo Rollouts (Kubernetes): بينفذ Canary Deployment (20% → 50% → 80% → 100%) مع Automated Gates.')
    add_paragraph(doc, 'الـ Traffic Router ده هو باب الدخول للمستخدم، بيختار إنها يروح لـ V1 ولا V2 حسب نسبة معينة، ويسجل كل prediction عشان نقدر نراقبها بعدين.')

    # 3. Tech Stack
    add_heading(doc, '3. الـ Tech Stack (التقنيات المستخدمة)', level=1)
    add_paragraph(doc, 'إحنا استخدمنا mix من تقنيات الـ ML, Serving, Monitoring, و Infrastructure:')
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'الفئة'
    hdr_cells[1].text = 'التقنيات'
    rows = [
        ('Model', 'Qwen2.5-1.5B (LLM), Transformers, PyTorch'),
        ('Quantization', 'GPTQ, AutoGPTQ, gptqmodel'),
        ('Experiment Tracking', 'MLflow'),
        ('Serving', 'FastAPI, vLLM, HTTPX'),
        ('Monitoring', 'Evidently AI, Prometheus, Grafana'),
        ('Orchestration', 'Kubernetes, Argo Rollouts'),
        ('Infrastructure', 'Docker, ConfigMaps, ServiceMonitors'),
    ]
    for r in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = r[0]
        row_cells[1].text = r[1]
    doc.add_paragraph()
    add_paragraph(doc, 'ليه اخترنا vLLM؟ لأنه optimized serving engine for LLMs، بيستخدم PagedAttention عشان يزيد الـ throughput ويقلل الـ latency.')
    add_paragraph(doc, 'ليه اخترنا MLflow؟ لأنه standard open-source لـ experiment tracking و model registry، بيسهل comparison بين runs وإدارة versions.')
    add_paragraph(doc, 'ليه اخترنا Evidently AI؟ لأنه بيقدم drift detection جاهز مع reports واضحة (HTML) وUI service.')

    # 4. Project Structure
    add_heading(doc, '4. هيكل المشروع (Project Structure)', level=1)
    add_paragraph(doc, 'المشروع مقسم لـ directories واضحة عشان يسهل الصيانة:')
    tree = """
ml-sentiment-app/
├── data/                       # الـ Datasets (train/eval)
├── training/
│   ├── configs/                # experiment_config.yaml (إعدادات MLflow و Model)
│   ├── src/
│   │   ├── train.py              # الـ Main script بتاع التدريب والتجارب
│   │   ├── data_loader.py        # تحميل وإنشاء datasets
│   │   ├── evaluate.py           # حساب الـ Metrics (Accuracy, F1, Precision, Recall)
│   │   ├── registration.py       # تسجيل الموديل في MLflow Registry
│   │   └── experiments/
│   │       └── few_shot.py       # تجارب Few-Shot Prompting
│   └── requirements.txt
├── quantization/
│   └── quantize_model.py         # GPTQ Quantization + Evaluation + MLflow Logging
├── serving/
│   ├── traffic_router.py         # FastAPI A/B Router
│   ├── Dockerfile
│   └── requirements.txt
├── monitoring/
│   ├── drift_detector.py         # Evidently Drift Detection
│   ├── send_predictions.py       # بيبعت test predictions لـ vLLM
│   ├── Dockerfile.analysis         # Docker image للـ Analysis Job
│   ├── analysis_entrypoint.sh    # Script بيشغل الـ Analysis Job
│   └── evidently_config.yaml     # إعدادات Evidently
├── k8s/                          # Kubernetes Manifests
│   ├── rollout.yaml              # Argo Rollouts Canary Config
│   ├── analysis-template.yaml    # Automated Drift Analysis Template
│   ├── service.yaml              # Stable + Canary Services
│   ├── configmap.yaml            # ConfigMaps (Env vars)
│   ├── grafana-dashboard.yaml    # Dashboard definition
│   └── pushgateway-servicemonitor.yaml
├── report/                       # صور ونتائج للـ Report
└── README.md
"""
    add_paragraph(doc, tree, indent=False)
    add_paragraph(doc, 'التقسيم ده مهم لأنه بيفصل concerns: كل directory مسؤول عن حاجة واحدة (training vs serving vs monitoring vs deployment). ده بيسهل onboarding لأي حد جديد في الفريق، وبيخلي الـ CI/CD pipeline أبسط.')

    # 5. Data
    add_heading(doc, '5. الـ Data (البيانات)', level=1)
    add_paragraph(doc, 'البيانات عبارة عن CSV files بسيطة:')
    add_bullet(doc, 'train_sentiment.csv: بيانات التدريب (لو احتجنا fine-tuning).')
    add_bullet(doc, 'eval_sentiment.csv: بيانات التقييم (Evaluation) — فيها text و label (positive/negative/neutral).')
    add_paragraph(doc, 'الـ Data Loader (data_loader.py) بيحمل الـ CSV ويحوله لـ list of dicts. وفي function create_sample_dataset() بتقدر تعمل dataset صغير بسرعة للـ testing.')
    add_paragraph(doc, 'ليه عملنا كده؟ عشان في المشروع ده مش بنعمل Traditional Fine-Tuning (لأننا بنستخدم LLM بـ Prompting)، فالبيانات محتاجة بس للـ Evaluation عشان نحسب الأداء.')

    # 6. Training & Few-Shot
    add_heading(doc, '6. Training & Few-Shot Prompting Experiments', level=1)
    add_paragraph(doc, 'في الـ train.py، إحنا بنعمل 2 حاجات رئيسية:')
    add_heading(doc, 'أ. Few-Shot Prompting Experiments', level=2)
    add_paragraph(doc, 'الفكرة: بدل ما ندرب الموديل من الصفر، بنعطيه examples قليلة (few-shot) في الـ prompt عشان "يتعلم" المهمة بسرعة.')
    add_paragraph(doc, 'few_shot.py بيعمل:')
    add_bullet(doc, '1-shot: prompt فيه example واحد (مثال positive/negative/neutral).')
    add_bullet(doc, '3-shot: 3 examples.')
    add_bullet(doc, '5-shot: 5 examples.')
    add_paragraph(doc, 'كل experiment بيحسب: Accuracy, F1-macro, Precision-macro, Recall-macro, Total Inference Time, Average Latency.')
    add_paragraph(doc, 'النتائج المشهورة: 1-shot (~65% accuracy), 3-shot (~72%), 5-shot (~75%). يعني زيادة الـ examples بتحسن الأداء لكن بتزود الـ latency شوية.')
    add_paragraph(doc, 'ليه بنعمل Few-Shot؟ لأن LLMs زي Qwen2.5-1.5B متدربة على tasks عامة، ومع الـ examples القليلة بتقدر تتكيف للـ task المحددة ( sentiment analysis) بدون fine-tuning مكلف.')
    add_heading(doc, 'ب. MLflow Integration', level=2)
    add_paragraph(doc, 'train.py بيستخدم MLflow عشان:')
    add_bullet(doc, 'log_param: model_name, n_shots, num_samples.')
    add_bullet(doc, 'log_metric: accuracy, f1_macro, precision_macro, recall_macro, latency.')
    add_bullet(doc, 'log_artifact: classification report (text) و confusion matrix (JSON).')
    add_paragraph(doc, 'الـ Experiment Name: sentiment-qwen2.5-experiments. الـ Tracking URI: http://localhost:5050.')
    add_paragraph(doc, 'ليه MLflow؟ عشان نقدر نcompare بين experiments بسهولة، نعرف أي run كان الأفضل، ونرجع لأي parameters أدت لأي results. ده مهم جداً في الـ team work والـ reproducibility.')

    # 7. Model Registry
    add_heading(doc, '7. Model Registry (MLflow)', level=1)
    add_paragraph(doc, 'registration.py بيسجل الموديل في MLflow Model Registry:')
    add_bullet(doc, 'بيبحث عن أفضل run (أعلى F1 score) لنوع معين من التجارب (few_shot أو quantization_eval).')
    add_bullet(doc, 'بيسجل الـ run كـ version جديد في registered model.')
    add_bullet(doc, 'بيطلق alias: champion (V1) للنسخة الأصلية، challenger (V2) للنسخة المضغوطة.')
    add_bullet(doc, 'بيضيف tag: vllm_model_path عشان نعرف المسار بتاع كل نسخة.')
    add_paragraph(doc, 'ده بيخلينا نعمل: A/B testing (champion vs challenger)، rollbacks (نرجع لـ champion لو الـ challenger فشل)، و versioning واضح.')
    add_paragraph(doc, 'في الـ MLflow UI، بنلاقي الموديل باسم: team-queens-a2-sentiment، مع 2 versions: Version 1 (champion) — V1، و Version 2 (challenger) — V2.')

    # 8. Quantization
    add_heading(doc, '8. Quantization (GPTQ 4-bit)', level=1)
    add_paragraph(doc, 'quantization/quantize_model.py بيعمل:')
    add_bullet(doc, '1. GPTQ Quantization: بيستخدم gptqmodel library عشان يضغط weights من 16-bit (FP16) لـ 4-bit. بيقلل model size من ~3.0 GB لـ ~1.0 GB (67% reduction).')
    add_bullet(doc, '2. Calibration Data: بيستخدم samples من wikitext dataset عشان يحسب quantization scales.')
    add_bullet(doc, '3. Post-Quantization Evaluation: بيحمل الـ model المضغوط وي evaluateه على eval dataset، ويحسب accuracy, f1, latency, model size.')
    add_bullet(doc, '4. MLflow Logging: يسجل كل metrics و parameters و artifacts (predictions CSV, confusion matrix PNG) لنفس الـ experiment.')
    add_paragraph(doc, 'ليه GPTQ؟ GPTQ (General-purpose Post-Training Quantization) هو quantization method بيضغط weights بتقسيمها لـ groups واستخدام approximate second-order information. بيحافظ على دقة الموديل بشكل ملحوظ (accuracy drop صغير: ~0.78 → ~0.76) بينما بيقلل الذاكرة والـ latency بشكل كبير.')
    add_paragraph(doc, 'الـ Results:')
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Light Grid Accent 1'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Model'
    hdr2[1].text = 'Accuracy'
    hdr2[2].text = 'F1'
    hdr2[3].text = 'Size'
    hdr2[4].text = 'Latency'
    qrows = [
        ('V1 (FP16)', '~0.78', '~0.76', '~3.0 GB', '~140 ms'),
        ('V2 (GPTQ 4-bit)', '~0.76', '~0.74', '~1.0 GB', '~90 ms'),
    ]
    for r in qrows:
        row_cells = table2.add_row().cells
        for i, v in enumerate(r):
            row_cells[i].text = v
    doc.add_paragraph()
    add_paragraph(doc, 'ليه عملنا quantization؟ في الـ Production، GPU memory غالي، والـ latency بيفرق مع الـ user. الـ 4-bit model بيقدر يشتغل على GPUs أضعف، وبيستجيب أسرع، وبنفس الـ quality تقريباً.')

    # 9. Serving & A/B Testing
    add_heading(doc, '9. Serving & A/B Testing (FastAPI Traffic Router)', level=1)
    add_paragraph(doc, 'serving/traffic_router.py عبارة عن FastAPI app:')
    add_heading(doc, 'أ. Endpoints', level=2)
    add_bullet(doc, '/health: Health check.')
    add_bullet(doc, '/predict: بيستقبل text (JSON) ويرجع {label, model_version, latency_ms}.')
    add_bullet(doc, '/logs: بيرجع prediction logs (JSONL) للـ monitoring.')
    add_bullet(doc, '/stats: بيرجع A/B statistics (counts, percentages, average latencies).')
    add_heading(doc, 'ب. Traffic Split', level=2)
    add_paragraph(doc, 'V1_WEIGHT = 0.80 (ممكن تتحكم فيه من environment variable). يعني 80% من الـ requests بيروحوا لـ V1 (original/champion) و 20% لـ V2 (quantized/challenger).')
    add_paragraph(doc, 'الـ Router بيستخدم random.random() لاختيار الـ version، وبعدين يبعت الـ request لـ vLLM server المناسب (V1_URL أو V2_URL) باستخدام httpx.AsyncClient.')
    add_heading(doc, 'ج. Prompt Template', level=2)
    add_paragraph(doc, 'الـ Prompt ثابت: "Classify the sentiment of the following text as positive, negative, or neutral.\\nText: {text}\\nSentiment:". ده بيضمن consistency بين V1 و V2.')
    add_heading(doc, 'د. Logging', level=2)
    add_paragraph(doc, 'كل prediction بيتسجل في JSONL file: timestamp, model_version, input_text, predicted_label, latency_ms. الـ logging ده أساسي لـ: Drift Detection, Debugging, Audit.')
    add_paragraph(doc, 'ليه FastAPI؟ لأنه سريع، asynchronous (بيستخدم async/await مع httpx)، وبيقدم automatic documentation (Swagger UI). وليه A/B Testing؟ عشان نختبر الـ V2 (quantized) في الـ Production بشكل آمن مع نسبة قليلة من الـ users، ولو نجح نزود النسبة.')

    # 10. Monitoring & Drift Detection
    add_heading(doc, '10. Monitoring & Drift Detection (Evidently AI + Prometheus)', level=1)
    add_paragraph(doc, 'monitoring/drift_detector.py بيعمل:')
    add_bullet(doc, 'بيقارن reference predictions (baseline) مع current predictions (canary/new).')
    add_bullet(doc, 'بيستخدم Evidently AI DataDriftPreset عشان يحسب drift في: text length, latency, prediction distribution.')
    add_bullet(doc, 'بيطلع HTML report (drift_report.html) ويوفره محلياً.')
    add_bullet(doc, 'لو --push مفعّل، بيبعت الـ report لـ Evidently UI Service (RemoteWorkspace).')
    add_bullet(doc, 'لو --pushgateway-url مفعّل، بيبعت drift_score كـ Prometheus metric (sentiment_drift_score) لـ Pushgateway.')
    add_paragraph(doc, 'send_predictions.py: بيبعت test texts لـ vLLM endpoint ويسجل الـ responses كـ JSONL، عشان نقدر نبني reference و current datasets.')
    add_paragraph(doc, 'evidently_config.yaml: بيحدد: check_interval (300s)، min_samples (50)، reference_window (500)، current_window (100)، numerical_method (Kolmogorov-Smirnov test)، categorical_method (Chi-squared test).')
    add_paragraph(doc, 'ليه Drift Detection مهمة؟ لأن في الـ Production، بيانات الـ users بتتغير مع الوقت (concept drift أو data drift). لو الموديل اتدرب على reviews عن electronics وجايله reviews عن clothes، الأداء هيقل. Evidently بيقلنا بدري إن في تغير.')
    add_paragraph(doc, 'ليه Prometheus؟ عشان يقدم time-series metrics لـ Grafana dashboards و Alertmanager. الـ Pushgateway مهم لأن الـ Analysis Job بيشتغل لفترة قصيرة (batch job) مش continuous service.')

    # 11. Kubernetes Deployment
    add_heading(doc, '11. Kubernetes Deployment (Canary with Argo Rollouts)', level=1)
    add_paragraph(doc, 'المشروع بيستخدم Kubernetes + Argo Rollouts عشان progressive delivery:')
    add_heading(doc, 'أ. Manifests في k8s/', level=2)
    add_bullet(doc, 'rollout.yaml: بيحدد الـ Rollout resource (بدل Deployment). فيه spec.strategy.canary بيحدد steps:')
    add_bullet(doc, '  20% traffic → Drift Analysis Gate (analysis-template)', level=1)
    add_bullet(doc, '  50% traffic → Pause 30s', level=1)
    add_bullet(doc, '  80% traffic → Pause 30s', level=1)
    add_bullet(doc, '  100% traffic → Rollout complete', level=1)
    add_bullet(doc, 'service.yaml: 2 Services (sentiment-model-stable و sentiment-model-canary). الـ Argo Rollouts بي manage الـ selector labels عشان يوجه traffic لـ stable أو canary pods.')
    add_bullet(doc, 'analysis-template.yaml: بيحدد AnalysisTemplate اسمه drift-analysis. بيشغل Job (analysis_job) بيستخدم image: registry.local:5000/sentiment-analysis-job:latest. الـ Job بيشغل send_predictions.py و drift_detector.py. وفي metric تانية: drift-score-check بتسأل Prometheus لو drift_score < 0.5. لو فشل، الـ Rollout بيتAbort.')
    add_bullet(doc, 'configmap.yaml: بيحتوي env vars: MLFLOW_TRACKING_URI, EVIDENTLY_URL, PROMETHEUS_PUSHGATEWAY_URL.')
    add_bullet(doc, 'grafana-dashboard.yaml: ConfigMap بيحتوي JSON dashboard definition للـ Drift Score.')
    add_bullet(doc, 'pushgateway-servicemonitor.yaml: ServiceMonitor عشان Prometheus يكتشف الـ Pushgateway.')
    add_heading(doc, 'ب. Flow بتاع Canary Deployment', level=2)
    add_paragraph(doc, '1. بنعمل kubectl apply -f k8s/.')
    add_paragraph(doc, '2. الـ Rollout بيحط 20% من الـ pods على النسخة الجديدة (V2/canary).')
    add_paragraph(doc, '3. الـ Analysis Template بيشغل Job: يبعت predictions لـ baseline و canary، يحسب drift، وي push metric لـ Prometheus.')
    add_paragraph(doc, '4. الـ drift-score-check بتسأل Prometheus: هل sentiment_drift_score < 0.5؟')
    add_paragraph(doc, '5. لو نعم: الـ Rollout بيكمل لـ 50% → 80% → 100%. لو لا: Abort ويرجع لـ 100% stable.')
    add_paragraph(doc, 'ليه Canary؟ عشان لو النسخة الجديدة فيها bug أو بتقل الأداء، بيأثر على 20% من الـ users بس، مش 100%. ده بيقلل الـ blast radius. وليه Automated Analysis Gates؟ عشان الـ decision يكون data-driven مش manual.')

    # 12. Docker
    add_heading(doc, '12. Docker & Containerization', level=1)
    add_paragraph(doc, 'serving/Dockerfile: بيستخدم multi-stage build:')
    add_bullet(doc, 'Stage 1 (builder): بيثبت dependencies في /install.')
    add_bullet(doc, 'Stage 2: بياخد /install بس، بيضيف traffic_router.py، بيشتغل بـ non-root user (appuser) للـ security.')
    add_paragraph(doc, 'monitoring/Dockerfile.analysis: بيحتوي evidently, pandas, scikit-learn, httpx, prometheus-client. بيستخدم analysis_entrypoint.sh كـ entrypoint.')
    add_paragraph(doc, 'ليه Multi-stage build؟ عشان image size يكون صغير (مش بنحتاج build tools في الـ final image). وليه Non-root user؟ security best practice في containers.')

    # 13. Results & Key Takeaways
    add_heading(doc, '13. النتائج والـ Key Takeaways', level=1)
    add_paragraph(doc, 'Few-Shot Performance:')
    table3 = doc.add_table(rows=1, cols=4)
    table3.style = 'Light Grid Accent 1'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = 'Experiment'
    hdr3[1].text = 'Accuracy'
    hdr3[2].text = 'F1 Score'
    hdr3[3].text = 'Latency (ms)'
    frows = [
        ('1-shot', '~0.65', '~0.62', '~120'),
        ('3-shot', '~0.72', '~0.70', '~145'),
        ('5-shot', '~0.75', '~0.73', '~160'),
    ]
    for r in frows:
        row_cells = table3.add_row().cells
        for i, v in enumerate(r):
            row_cells[i].text = v
    doc.add_paragraph()
    add_paragraph(doc, 'Quantization vs Baseline:')
    table4 = doc.add_table(rows=1, cols=5)
    table4.style = 'Light Grid Accent 1'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = 'Model'
    hdr4[1].text = 'Accuracy'
    hdr4[2].text = 'F1'
    hdr4[3].text = 'Size'
    hdr4[4].text = 'Latency'
    qrows2 = [
        ('V1 (FP16)', '~0.78', '~0.76', '~3.0 GB', '~140 ms'),
        ('V2 (GPTQ 4-bit)', '~0.76', '~0.74', '~1.0 GB', '~90 ms'),
    ]
    for r in qrows2:
        row_cells = table4.add_row().cells
        for i, v in enumerate(r):
            row_cells[i].text = v
    doc.add_paragraph()
    add_paragraph(doc, 'Key Takeaways:')
    add_bullet(doc, 'Few-Shot Prompting بيحسن الأداء بشكل ملحوظ بدون fine-tuning.')
    add_bullet(doc, 'GPTQ 4-bit Quantization بيقلل size بـ 67% و latency بـ ~35% مع drop صغير في الـ accuracy.')
    add_bullet(doc, 'A/B Testing + Traffic Routing بيسمح بـ safe rollout للـ models الجديدة.')
    add_bullet(doc, 'Drift Detection + Prometheus + Grafana = Monitoring شبه real-time.')
    add_bullet(doc, 'Canary Deployments + Automated Gates = Deployment بأمان و data-driven.')
    add_bullet(doc, 'MLflow Tracking + Registry = Reproducibility و Version Management.')

    # 14. Interview Questions
    add_heading(doc, '14. أسئلة متوقعة في الـ Interview وإزاي تجاوب عليها', level=1)
    add_paragraph(doc, 'س1: "إيه الفرق بين Few-Shot Prompting و Fine-Tuning؟"')
    add_paragraph(doc, 'ج: Few-Shot بيستخدم examples في الـ prompt بدون تغيير weights الموديل. سريع، مش محتاج GPU قوي، ومناسب لـ tasks بسيطة. Fine-Tuning بيغير weights، محتاج data أكتر، GPU أقوى، وقت أطول، لكن بيحسن الأداء أكتر على tasks معقدة. في المشروع، اخترنا Few-Shot لأن Sentiment Analysis task بسيطة والـ model صغير (1.5B parameters).')
    add_paragraph(doc, 'س2: "ليه استخدمتو GPTQ مش INT8 أو AWQ؟"')
    add_paragraph(doc, 'ج: GPTQ بيقدم compression أعلى (4-bit) مع دقة محتفظة بشكل كويس. INT8 بيحافظ أكتر على الدقة بس compression أقل. AWQ جديد وكويس بس GPTQ أشهر وdocumentation أحسن للـ Qwen models. اخترنا GPTQ لأنه balance بين size reduction و accuracy.')
    add_paragraph(doc, 'س3: "إزاي بيشتغل الـ Drift Detection؟"')
    add_paragraph(doc, 'ج: بيستخدم Evidently AI عشان يقارن reference dataset (baseline) مع current dataset (canary). بيحسب drift score للـ numerical features (KS test) و categorical features (Chi-squared). لو الـ score > 0.5، يعتبر drift detected. الـ score بيتبعت لـ Prometheus Pushgateway، وGrafana بيظهره، وArgo Rollouts بيستخدمه كـ gate.')
    add_paragraph(doc, 'س4: "إيه الفرق بين Blue-Green و Canary Deployment؟"')
    add_paragraph(doc, 'ج: Blue-Green: بيشغل 2 environments كاملين (blue=old, green=new) وبي switch الـ traffic دفعة واحدة. Canary: بيودي نسبة صغيرة من الـ traffic للنسخة الجديدة (مثلاً 20%) ويزودها تدريجياً. Canary أأمن لأن لو في مشكلة بيأثر على قليل من الـ users. في المشروع، استخدمنا Canary مع Argo Rollouts.')
    add_paragraph(doc, 'س5: "ليه استخدمتو Prometheus Pushgateway مش normal scraping؟"')
    add_paragraph(doc, 'ج: لأن الـ Analysis Job بيشتغل كـ batch job (Kubernetes Job) مش continuous service. الـ Prometheus normal scraping بيحتاج endpoint شغال دايماً. Pushgateway بيسمح للـ batch jobs إنها تبعت metrics مرة واحدة.')
    add_paragraph(doc, 'س6: "إزاي بيشتغل الـ Traffic Router؟"')
    add_paragraph(doc, 'ج: FastAPI app بيستقبل request على /predict. بيختار version (V1 أو V2) حسب weight (80/20). بيبعت الـ request لـ vLLM server المناسب. بيستقبل الـ response، بي parse الـ sentiment label، بي سجل في JSONL، وبيرجع للـ client.')
    add_paragraph(doc, 'س7: "إزاي بيشتغل الـ MLflow Model Registry؟"')
    add_paragraph(doc, 'ج: registration.py بيبحث عن أفضل run (أعلى F1) في experiment. بيسجل الـ run كـ version في registered model. بيطلق alias (champion/challenger). بيضيف tag (vllm_model_path) عشان نعرف الموديل الفعلي. ده بيسهل A/B testing والـ rollback.')

    # 15. Summary
    add_heading(doc, '15. ملخص سريع (Executive Summary)', level=1)
    add_paragraph(doc, 'المشروع بيظهر End-to-End MLOps Pipeline لـ Sentiment Analysis LLM:')
    add_bullet(doc, 'بنستخدم Qwen2.5-1.5B مع Few-Shot Prompting (1/3/5-shot) ون track بالـ MLflow.')
    add_bullet(doc, 'بنضغط الموديل بـ GPTQ 4-bit (3x smaller) ون evaluate الأداء.')
    add_bullet(doc, 'بنقدم API بـ FastAPI Traffic Router بـ A/B Testing (80/20 split).')
    add_bullet(doc, 'بنراقب Data Drift بـ Evidently AI و Prometheus.')
    add_bullet(doc, 'بننشر بـ Kubernetes Canary Deployments باستخدام Argo Rollouts + Automated Gates.')
    add_paragraph(doc, 'الـ Goal: LLM في Production بأمان، مراقبة، optimization، وdeployment تدريجي.')

    doc.save('C:\\Users\\Admin\\Downloads\\devops\\ml-sentiment-app-a2\\Interview_Guide_ML_Sentiment_MLOps.docx')
    print('Done! File saved at: C:\\Users\\Admin\\Downloads\\devops\\ml-sentiment-app-a2\\Interview_Guide_ML_Sentiment_MLOps.docx')

if __name__ == '__main__':
    main()
